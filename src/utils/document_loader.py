# scr/utils/document_loader.py
import logging
import re
from pathlib import Path
from typing import List, Dict, Optional, Tuple
from dataclasses import dataclass

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    from docx import Document
except ImportError:
    Document = None

logger = logging.getLogger(__name__)

@dataclass
class DocumentChunk:
    """Структурированный чанк документа с метаданными"""
    text: str
    source: str
    doc_type: str
    section: Optional[str] = None
    clause: Optional[str] = None
    page: Optional[int] = None
    full_path: Optional[str] = None
    
    def to_dict(self) -> Dict:
        return {
            "text": self.text,
            "source": self.source,
            "doc_type": self.doc_type,
            "section": self.section,
            "clause": self.clause,
            "page": self.page,
            "full_path": self.full_path
        }


class DocumentParser:
    """Парсер нормативных документов с извлечением структуры"""
    
    # Паттерны для распознавания структуры ГОСТ и других стандартов
    SECTION_PATTERNS = [
        r'^(\d+)\s+([А-Я][а-я]+(?:\s+[А-Я][а-я]+)*)\s*$',  # "1 ОБЩИЕ ПОЛОЖЕНИЯ"
        r'^РАЗДЕЛ\s+(\d+(?:\.\d+)?)\s*[:\-]?\s*(.+)$',      # "РАЗДЕЛ 5: Требования"
        r'^ГЛАВА\s+(\d+(?:\.\d+)?)\s*[:\-]?\s*(.+)$',       # "ГЛАВА 3. Контроль"
    ]
    
    CLAUSE_PATTERNS = [
        r'^(\d+(?:\.\d+){1,3})\.\s+(.+)$',                  # "6.4.2. Требования к..."
        r'^п\.?\s*(\d+(?:\.\d+){1,3})\s*[:\-]?\s*(.+)$',   # "п. 6.4.2: Требования"
        r'^(\d+(?:\.\d+){1,2})\s+([А-Я][а-я].+)$',         # "5.2 Требования безопасности"
    ]
    
    def __init__(self):
        self.current_section = None
        self.current_clause = None
    
    def _parse_line_structure(self, line: str) -> Tuple[Optional[str], Optional[str], str]:
        """
        Анализ строки на наличие номера раздела/пункта
        
        Returns:
            (section_number, clause_number, clean_text)
        """
        line = line.strip()
        
        # Проверяем разделы
        for pattern in self.SECTION_PATTERNS:
            match = re.match(pattern, line, re.IGNORECASE)
            if match:
                self.current_section = match.group(1)
                self.current_clause = None
                return self.current_section, None, line
        
        # Проверяем пункты
        for pattern in self.CLAUSE_PATTERNS:
            match = re.match(pattern, line)
            if match:
                clause_num = match.group(1)
                # Определяем, это подраздел раздела или отдельный пункт
                if clause_num.count('.') == 0 and self.current_section:
                    # Это подраздел текущего раздела (например, "5.1" когда раздел "5")
                    self.current_clause = clause_num
                else:
                    self.current_clause = clause_num
                return self.current_section, self.current_clause, line
        
        return self.current_section, self.current_clause, line
    
    def parse_text_by_structure(self, text: str, source: str, doc_type: str) -> List[DocumentChunk]:
        """
        Разбивка текста на структурированные чанки с сохранением иерархии
        
        Args:
            text: Полный текст документа
            source: Имя файла
            doc_type: Тип документа
        
        Returns:
            Список DocumentChunk с метаданными
        """
        chunks = []
        self.current_section = None
        self.current_clause = None
        
        lines = text.split('\n')
        current_content = []
        current_page = 1
        
        for line in lines:
            # Проверяем, не начинается ли новая структурная единица
            section, clause, clean_line = self._parse_line_structure(line)
            
            # Если нашли новый раздел или пункт, сохраняем предыдущий чанк
            if (section != self.current_section or clause != self.current_clause) and current_content:
                chunk_text = '\n'.join(current_content).strip()
                if chunk_text and len(chunk_text) > 50:  # Пропускаем слишком короткие
                    chunks.append(DocumentChunk(
                        text=chunk_text,
                        source=source,
                        doc_type=doc_type,
                        section=self.current_section,
                        clause=self.current_clause,
                        page=current_page,
                        full_path=source
                    ))
                current_content = []
            
            # Обновляем текущую структуру
            if section:
                self.current_section = section
            if clause:
                self.current_clause = clause
            
            # Добавляем строку в текущий контент
            if clean_line:
                current_content.append(clean_line)
            
            # Простая эвристика для определения страниц (каждые 40 строк)
            if len(current_content) % 40 == 0:
                current_page += 1
        
        # Сохраняем последний чанк
        if current_content:
            chunk_text = '\n'.join(current_content).strip()
            if chunk_text and len(chunk_text) > 50:
                chunks.append(DocumentChunk(
                    text=chunk_text,
                    source=source,
                    doc_type=doc_type,
                    section=self.current_section,
                    clause=self.current_clause,
                    page=current_page,
                    full_path=source
                ))
        
        logger.info(f"Разбито на {len(chunks)} структурированных чанков")
        return chunks


def _extract_pdf_with_structure(file_path: str) -> List[DocumentChunk]:
    """Извлечение текста из PDF с сохранением структуры"""
    if PyPDF2 is None:
        raise ImportError("Установите PyPDF2: pip install PyPDF2")
    
    parser = DocumentParser()
    chunks = []
    
    with open(file_path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        full_text = ""
        
        for page_num, page in enumerate(reader.pages, 1):
            page_text = page.extract_text() or ""
            full_text += page_text + "\n\n"
        
        # Парсим структуру
        source_name = Path(file_path).name
        chunks = parser.parse_text_by_structure(full_text, source_name, "pdf")
        
        # Добавляем информацию о страницах
        for chunk in chunks:
            chunk.full_path = file_path
    
    return chunks


def _extract_docx_with_structure(file_path: str) -> List[DocumentChunk]:
    """Извлечение текста из DOCX с сохранением структуры"""
    if Document is None:
        raise ImportError("Установите python-docx: pip install python-docx")
    
    parser = DocumentParser()
    doc = Document(file_path)
    
    # Собираем текст с учетом стилей
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            style = para.style.name if para.style else "Normal"
            paragraphs.append({
                "text": para.text.strip(),
                "style": style
            })
    
    full_text = "\n\n".join([p["text"] for p in paragraphs])
    source_name = Path(file_path).name
    
    chunks = parser.parse_text_by_structure(full_text, source_name, "docx")
    
    for chunk in chunks:
        chunk.full_path = file_path
    
    return chunks


def load_documents(directory: str, use_structure: bool = True) -> List[Dict]:
    """
    Загружает все поддерживаемые документы из указанной директории (рекурсивно).
    
    Args:
        directory: Путь к директории
        use_structure: Если True, пытается извлечь структуру (разделы, пункты)
    
    Returns:
        Список словарей с документами и метаданными
    """
    docs = []
    dir_path = Path(directory)

    if not dir_path.exists():
        raise FileNotFoundError(f"Директория не найдена: {directory}")

    logger.info(f"Загрузка документов из {directory}...")

    for file_path in dir_path.rglob("*"):
        if not file_path.is_file():
            continue

        ext = file_path.suffix.lower()
        doc_type = ext.lstrip(".")
        
        try:
            if ext == ".pdf":
                if use_structure:
                    chunks = _extract_pdf_with_structure(str(file_path))
                else:
                    text = _extract_pdf_simple(str(file_path))
                    chunks = [DocumentChunk(text=text, source=file_path.name, doc_type=doc_type, full_path=str(file_path))]
                    
            elif ext == ".docx":
                if use_structure:
                    chunks = _extract_docx_with_structure(str(file_path))
                else:
                    text = _extract_docx_simple(str(file_path))
                    chunks = [DocumentChunk(text=text, source=file_path.name, doc_type=doc_type, full_path=str(file_path))]
                    
            elif ext == ".txt":
                text = file_path.read_text(encoding="utf-8")
                chunks = [DocumentChunk(text=text, source=file_path.name, doc_type=doc_type, full_path=str(file_path))]
            else:
                continue

            # Фильтруем пустые чанки
            valid_chunks = [c for c in chunks if c.text.strip()]
            
            if valid_chunks:
                docs.extend([c.to_dict() for c in valid_chunks])
                logger.debug(f"  {file_path.name}: {len(valid_chunks)} чанков")
            else:
                logger.warning(f"  {file_path.name}: документ пуст или не распарсен")
                
        except Exception as e:
            logger.error(f"Ошибка чтения {file_path.name}: {e}", exc_info=True)

    logger.info(f"✓ Загружено {len(docs)} документов из {directory}")
    return docs


# Простые версии для обратной совместимости
def _extract_pdf_simple(file_path: str) -> str:
    if PyPDF2 is None:
        raise ImportError("Установите PyPDF2: pip install PyPDF2")
    text = ""
    with open(file_path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text


def _extract_docx_simple(file_path: str) -> str:
    if Document is None:
        raise ImportError("Установите python-docx: pip install python-docx")
    doc = Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])


# Для обратной совместимости
def load_documents_simple(directory: str) -> List[Dict]:
    """Старая версия без извлечения структуры"""
    return load_documents(directory, use_structure=False)