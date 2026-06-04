# scripts\index_examples.py
"""
Индексация примеров документов в RAG-базу
Поддержка: PDF, DOCX, DOC, MD
"""
import os
import json
import logging
import re
from pathlib import Path
from typing import Dict, List, Any, Optional

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


def extract_text_from_pdf(pdf_path: str) -> Optional[str]:
    """Извлечение текста из PDF"""
    try:
        import PyPDF2
        
        text_content = []
        with open(pdf_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_content.append(text)
        
        return "\n".join(text_content) if text_content else None
    except Exception as e:
        logger.error(f"Ошибка чтения PDF {pdf_path}: {e}")
        return None


def extract_text_from_docx(docx_path: str) -> Optional[str]:
    """Извлечение текста из DOCX"""
    try:
        from docx import Document
        
        doc = Document(docx_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n".join(paragraphs) if paragraphs else None
    except Exception as e:
        logger.error(f"Ошибка чтения DOCX {docx_path}: {e}")
        return None


def extract_text_from_doc(doc_path: str) -> Optional[str]:
    """Извлечение текста из старого формата DOC"""
    try:
        # Попытка 1: antiword (лучшее качество)
        import subprocess
        result = subprocess.run(
            ['antiword', doc_path],
            capture_output=True,
            text=True,
            timeout=30
        )
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout
        
        # Попытка 2: python-docx (может работать с некоторыми DOC)
        from docx import Document
        doc = Document(doc_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n".join(paragraphs) if paragraphs else None
        
    except Exception as e:
        logger.warning(f"Не удалось прочитать DOC {doc_path}: {e}")
        return None


def extract_text_from_markdown(md_path: str) -> Optional[str]:
    """Извлечение текста из Markdown"""
    try:
        with open(md_path, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        logger.error(f"Ошибка чтения Markdown {md_path}: {e}")
        return None


def extract_structure_from_text(text: str, source_file: str) -> Optional[Dict[str, Any]]:
    """Улучшенное извлечение структуры с поддержкой вложенности и корректной нумерацией"""
    if not text or len(text.strip()) < 50:
        logger.debug(f"Текст слишком короткий для извлечения структуры: {len(text) if text else 0} символов")
        return None
    
    structure = {"root": [], "metadata": {"source_file": source_file}}
    lines = text.split('\n')
    
    # Стек для отслеживания вложенности: [(уровень, раздел, номер_родителя)]
    stack: List[tuple] = []
    section_counter = {"1": 0}  # Счётчик для генерации номеров
    
    logger.debug(f"Начало парсинга структуры для {source_file}, всего строк: {len(lines)}")
    
    for line_idx, line in enumerate(lines):
        line_stripped = line.strip()
        
        # Пропускаем пустые строки и слишком короткие
        if not line_stripped or len(line_stripped) < 3:
            continue
        
        # Пропускаем строки, которые явно не являются заголовками
        if line_stripped.startswith(('—', '•', '*', '1.', '2.', '3.')) and not re.match(r'^\d+[\.\)]\s+[А-Я]', line_stripped):
            continue
        
        level = None
        title = None
        number = None
        
        # === Паттерн 1: # Заголовок (Markdown) ===
        md_match = re.match(r'^(#{1,6})\s+(.+)$', line_stripped)
        if md_match:
            level = len(md_match.group(1))
            title = md_match.group(2).strip()
            number = str(level)  # Временный номер для markdown
        
        # === Паттерн 2: 1.2.3 Заголовок (нумерованный) ===
        if not md_match:
            # Более точный паттерн: номер + разделитель + заголовок на заглавную
            num_match = re.match(r'^(\d+(?:\.\d+){0,4})[\.\s:—\-]+\s*([А-ЯЁ].+)$', line_stripped)
            if num_match:
                number = num_match.group(1)
                title = num_match.group(2).strip()
                level = number.count('.') + 1
                logger.debug(f"Строка {line_idx}: найден нумерованный заголовок '{title}' (уровень {level}, номер {number})")
        
        # === Паттерн 3: Ключевые слова без номера ===
        if not level and not title:
            important = [
                "общие положения", "термины", "нормативные", "требования", 
                "ответственность", "контроль", "заключ", "цели", "задачи",
                "область применения", "принципы"
            ]
            line_lower = line_stripped.lower()
            if any(kw in line_lower for kw in important):
                # Проверяем, что это действительно заголовок (не часть текста)
                if len(line_stripped) < 100 and line_stripped[0].isupper():
                    level = 1
                    title = line_stripped
                    # Извлекаем номер если есть в начале
                    num_prefix = re.match(r'^(\d+(?:\.\d+)*)[\.\s]+', line_stripped)
                    if num_prefix:
                        number = num_prefix.group(1)
                    logger.debug(f"Строка {line_idx}: найден ключевой заголовок '{title}'")
        
        # === Обработка найденного раздела ===
        if level and title:
            # Генерируем номер если не определён
            if not number:
                parent_num = stack[-1][2] if stack else "0"
                section_counter[parent_num] = section_counter.get(parent_num, 0) + 1
                number = f"{parent_num}.{section_counter[parent_num]}" if parent_num != "0" else str(section_counter[parent_num])
            
            section = {
                "title": title,
                "number": number,
                "level": level,
                "children": [],
                "content_preview": "",
                "line_index": line_idx  # Для отладки
            }
            
            # === Корректная вложенность через стек ===
            # Удаляем из стека все элементы с уровнем >= текущего
            while stack and stack[-1][0] >= level:
                stack.pop()
            
            if not stack:
                # Уровень 1 — добавляем в корень
                structure["root"].append(section)
                stack.append((level, section, number))
                logger.debug(f"  → Добавлен в корень: {number} '{title}'")
            else:
                # Вложенный раздел — добавляем в детей последнего родителя
                parent = stack[-1][1]
                parent["children"].append(section)
                stack.append((level, section, number))
                logger.debug(f"  → Добавлен как ребёнок {stack[-2][2]}: {number} '{title}'")
    
    # Пост-обработка
    if not structure["root"]:
        logger.warning(f"Не найдено ни одного раздела в {source_file}")
        return None
    
    structure = _ensure_mandatory_sections(structure)
    structure["metadata"]["sections_count"] = _count_all_sections(structure["root"])
    
    logger.info(f"✓ Извлечена структура: {len(structure['root'])} корневых разделов, "
               f"{structure['metadata']['sections_count']} всего")
    
    return structure

def _count_all_sections(sections: List[Dict]) -> int:
    """Рекурсивный подсчёт всех разделов включая вложенные"""
    count = len(sections)
    for section in sections:
        if section.get("children"):
            count += _count_all_sections(section["children"])
    return count

def _ensure_mandatory_sections(structure: Dict, min_count: int = 8) -> Dict:
    """Гарантирует наличие обязательных разделов в НАЧАЛЕ структуры"""
    mandatory_first = [
        {"title": "Общие положения", "level": 1},
        {"title": "Нормативные ссылки", "level": 1}, 
        {"title": "Термины и определения", "level": 1},
    ]
    
    mandatory_any = [
        {"title": "Требования к защите", "level": 1},
        {"title": "Ответственность", "level": 1},
        {"title": "Контроль и аудит", "level": 1},
        {"title": "Заключительные положения", "level": 1},
    ]
    
    existing = [s["title"].lower() for s in structure["root"]]
    root = structure["root"]
    
    # 1. Добавляем обязательные ПЕРВЫЕ разделы в начало (если нет)
    for req in reversed(mandatory_first):  # reversed для сохранения порядка при insert(0)
        if not any(req["title"].lower() in t for t in existing):
            root.insert(0, {
                "title": req["title"],
                "number": str(len(root) + 1),
                "level": req["level"],
                "children": [],
                "content_preview": f"Раздел '{req['title']}' добавлен автоматически"
            })
    
    # 2. Добавляем остальные обязательные разделы в конец (если нет)
    for req in mandatory_any:
        if not any(req["title"].lower() in t for t in existing):
            root.append({
                "title": req["title"], 
                "number": str(len(root) + 1),
                "level": req["level"],
                "children": [],
                "content_preview": f"Раздел '{req['title']}' добавлен автоматически"
            })
    
    # 3. Перенумеровываем разделы последовательно
    for i, section in enumerate(root, 1):
        section["number"] = str(i)
    
    structure["root"] = root
    return structure

def determine_doc_type(filename: str) -> str:
    """Определение типа документа по имени и содержимому"""
    filename_lower = filename.lower()
    
    if any(word in filename_lower for word in ['политика', 'policy']):
        return "policy"
    elif any(word in filename_lower for word in ['инструкция', 'instruk', 'инстр']):
        return "instruction"
    elif any(word in filename_lower for word in ['регламент', 'reglament']):
        return "regulation"
    elif any(word in filename_lower for word in ['угроз', 'threat', 'model']):
        return "threat_model"
    elif any(word in filename_lower for word in ['риск', 'risk']):
        return "risk_assessment"
    elif any(word in filename_lower for word in ['доступ', 'access', 'control']):
        return "access_control"
    else:
        return "policy"  # По умолчанию

def index_examples(examples_dir: str = "data/documents/examples"):
    """
    Индексирует все примеры документов из папки.
    Поддерживаемые форматы: PDF, DOCX, DOC, MD
    """
    from src.services.rag_service import RAGService
    rag = RAGService()
    
    examples_path = Path(examples_dir)
    
    if not examples_path.exists():
        logger.error(f"Папка примеров не найдена: {examples_path}")
        return
    
    # Словарь для подсчета статистики
    stats = {
        "total_files": 0,
        "indexed_files": 0,
        "failed_files": [],
        "by_format": {"pdf": 0, "docx": 0, "doc": 0, "md": 0}
    }
    
    logger.info(f"Сканирование папки: {examples_path.absolute()}")
    
    # Обрабатываем все файлы в папке
    for file_path in examples_path.iterdir():
        if file_path.is_file():
            stats["total_files"] += 1
            filename = file_path.name
            suffix = file_path.suffix.lower()
            
            logger.info(f"\nОбработка файла #{stats['total_files']}: {filename}")
            
            # Извлечение текста в зависимости от формата
            text = None
            if suffix == '.pdf':
                logger.info("   Формат: PDF")
                text = extract_text_from_pdf(str(file_path))
                stats["by_format"]["pdf"] += 1
            elif suffix == '.docx':
                logger.info("   Формат: DOCX")
                text = extract_text_from_docx(str(file_path))
                stats["by_format"]["docx"] += 1
            elif suffix == '.doc':
                logger.info("   Формат: DOC (старый)")
                text = extract_text_from_doc(str(file_path))
                stats["by_format"]["doc"] += 1
            elif suffix in ['.md', '.markdown']:
                logger.info("   Формат: Markdown")
                text = extract_text_from_markdown(str(file_path))
                stats["by_format"]["md"] += 1
            else:
                logger.warning(f"   Неподдерживаемый формат: {suffix}")
                stats["failed_files"].append((filename, "Неподдерживаемый формат"))
                continue
            
            if not text:
                logger.warning(f"   Не удалось извлечь текст")
                stats["failed_files"].append((filename, "Пустой текст"))
                continue
            
            logger.info(f"   ✓ Извлечено {len(text)} символов")
            
            # Извлечение структуры
            structure = extract_structure_from_text(text, str(file_path))
            
            if not structure or not structure["root"]:
                logger.warning(f"   Не удалось извлечь структуру")
                stats["failed_files"].append((filename, "Нет структуры"))
                continue
            
            # Определение типа документа
            doc_type = determine_doc_type(filename)
            structure["metadata"]["doc_type"] = doc_type
            
            logger.info(f"   ✓ Найдено разделов: {len(structure['root'])}")
            logger.info(f"   Тип документа: {doc_type}")
            
            # Добавление в RAG
            try:
                
                rag.add_example_structure(
                    structure=structure,  # ← Передаём саму структуру, не JSON-строку
                    source_file=str(file_path),
                    doc_type=doc_type,
                    organization=file_path.stem,
                    industry="",  # Можно определить по названию организации
                    regime=""     # Можно определить по профилю
                )
                
                stats["indexed_files"] += 1
                logger.info(f"   ✓ Проиндексирован успешно")
                
            except Exception as e:
                logger.error(f"   Ошибка добавления в RAG: {e}")
                stats["failed_files"].append((filename, str(e)))
    
    # Итоговая статистика
    logger.info("\n" + "="*60)
    logger.info("ИТОГИ ИНДЕКСАЦИИ ПРИМЕРОВ")
    logger.info("="*60)
    logger.info(f"Всего файлов: {stats['total_files']}")
    logger.info(f"Проиндексировано: {stats['indexed_files']}")
    logger.info(f"Не удалось: {len(stats['failed_files'])}")
    logger.info(f"\nПо форматам:")
    for fmt, count in stats["by_format"].items():
        logger.info(f"   {fmt.upper()}: {count}")
    
    if stats["failed_files"]:
        logger.warning(f"\nОшибки:")
        for filename, error in stats["failed_files"]:
            logger.warning(f"   {filename}: {error}")
    
    # Общая статистика RAG
    rag_stats = rag.get_statistics()
    logger.info(f"\nОбщая статистика RAG: {rag_stats}")
    logger.info("="*60)


if __name__ == "__main__":
    index_examples()