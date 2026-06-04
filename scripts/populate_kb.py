# scripts/populate_kb.py
import os
import sys
import re
import json
import logging
from pathlib import Path
from typing import List, Dict, Optional, Tuple
from datetime import datetime

# Добавляем корень проекта в путь
sys.path.append(str(Path(__file__).parent.parent))

from src.database import get_chroma_collection, get_neo4j_driver, get_embedding_model

# Настройка логирования
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s | %(levelname)-8s | %(message)s',
    handlers=[
        logging.StreamHandler(sys.stdout),
        logging.FileHandler('logs/populate_kb.log', encoding='utf-8')
    ]
)
logger = logging.getLogger(__name__)


class DocumentParser:
    """Универсальный парсер документов различных форматов"""
    
    @staticmethod
    def parse_pdf(file_path: Path) -> Tuple[str, Dict]:
        """
        Извлечение текста из PDF с метаданными.
        """
        logger.info(f"Парсинг PDF: {file_path.name}")
        
        try:
            import pdfplumber
            
            metadata = {
                "source_file": file_path.name,
                "format": "pdf",
                "pages": 0,
                "extracted_at": datetime.now().isoformat()
            }
            
            full_text = ""
            with pdfplumber.open(file_path) as pdf:
                metadata["pages"] = len(pdf.pages)
                for i, page in enumerate(pdf.pages, 1):
                    text = page.extract_text()
                    if text:
                        full_text += f"\n[Страница {i}]\n{text}"
            
            return full_text, metadata
            
        except Exception as e:
            logger.error(f"Ошибка парсинга PDF {file_path.name}: {e}")
            return "", {}
    
    @staticmethod
    def parse_docx(file_path: Path) -> Tuple[str, Dict]:
        """
        Извлечение текста из DOCX.
        """
        logger.info(f"Парсинг DOCX: {file_path.name}")
        
        try:
            from docx import Document
            
            doc = Document(file_path)
            
            metadata = {
                "source_file": file_path.name,
                "format": "docx",
                "paragraphs": len(doc.paragraphs),
                "extracted_at": datetime.now().isoformat()
            }
            
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            return "\n\n".join(text_parts), metadata
            
        except Exception as e:
            logger.error(f"Ошибка парсинга DOCX {file_path.name}: {e}")
            return "", {}
    
    @staticmethod
    def parse_odt(file_path: Path) -> Tuple[str, Dict]:
        """
        Извлечение текста из ODT (OpenDocument Text).
        """
        logger.info(f"Парсинг ODT: {file_path.name}")
        
        try:
            from odf.opendocument import load
            from odf.text import P
            
            doc = load(str(file_path))
            
            metadata = {
                "source_file": file_path.name,
                "format": "odt",
                "extracted_at": datetime.now().isoformat()
            }
            
            text_parts = []
            for paragraph in doc.getElementsByType(P):
                if paragraph.textContent.strip():
                    text_parts.append(paragraph.textContent)
            
            return "\n\n".join(text_parts), metadata
            
        except ImportError:
            logger.warning("Установите odfpy для поддержки ODT: pip install odfpy")
            return "", {}
        except Exception as e:
            logger.error(f"Ошибка парсинга ODT {file_path.name}: {e}")
            return "", {}
    
    @staticmethod
    def parse_text(file_path: Path) -> Tuple[str, Dict]:
        """
        Извлечение текста из TXT/MD файлов.
        """
        logger.info(f"Чтение текста: {file_path.name}")
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            metadata = {
                "source_file": file_path.name,
                "format": file_path.suffix[1:].lower(),
                "size_bytes": len(content.encode('utf-8')),
                "extracted_at": datetime.now().isoformat()
            }
            
            return content, metadata
            
        except Exception as e:
            logger.error(f"Ошибка чтения {file_path.name}: {e}")
            return "", {}


class TextChunker:
    """Разбиение текста на чанки с сохранением структуры"""
    
    @staticmethod
    def chunk_regulatory_document(text: str, source_metadata: Dict) -> List[Dict]:
        """
        Разбиение нормативного документа на смысловые разделы.
        """
        chunks = []
        
        # Паттерны для поиска разделов
        section_patterns = [
            r'(\d+\.\d+\.\d+\s+[\u0410-\u042F][\u0430-\u044F]+)',  # 1.1.1 Заголовок
            r'(\d+\.\d+\s+[\u0410-\u042F][\u0430-\u044F]+)',        # 1.1 Заголовок
            r'(\d+\s+[\u0410-\u042F][\u0430-\u044F]+)',             # 1 Заголовок
            r'(ПРИЛОЖЕНИЕ\s+[\u0410-\u042F])',                      # ПРИЛОЖЕНИЕ А
            r'(ГОСТ\s+\d)',                                          # ГОСТ Р 57580
        ]
        
        # Объединяем паттерны
        combined_pattern = '|'.join(f'({p})' for p in section_patterns)
        
        # Разбиваем текст
        sections = re.split(combined_pattern, text, flags=re.IGNORECASE)
        
        section_num = 0
        current_section = ""
        current_title = ""
        
        for part in sections:
            if not part or not part.strip():
                continue
            
            # Проверяем, является ли часть заголовком
            is_header = any(
                re.match(pattern.strip('()'), part.strip(), re.IGNORECASE)
                for pattern in section_patterns
            )
            
            if is_header:
                # Сохраняем предыдущий раздел
                if current_section.strip() and len(current_section) > 100:
                    section_num += 1
                    chunks.append({
                        "id": f"{source_metadata['source_file']}_section_{section_num}",
                        "text": current_section.strip(),
                        "section_number": current_title.strip() if current_title else f"Раздел {section_num}",
                        "source_file": source_metadata["source_file"],
                        "doc_type": "regulatory",
                        "format": source_metadata.get("format", "unknown"),
                        "page": section_num
                    })
                
                current_title = part.strip()
                current_section = part.strip() + "\n"
            else:
                current_section += part + "\n"
        
        # Добавляем последний раздел
        if current_section.strip() and len(current_section) > 100:
            section_num += 1
            chunks.append({
                "id": f"{source_metadata['source_file']}_section_{section_num}",
                "text": current_section.strip(),
                "section_number": current_title.strip() if current_title else f"Раздел {section_num}",
                "source_file": source_metadata["source_file"],
                "doc_type": "regulatory",
                "format": source_metadata.get("format", "unknown"),
                "page": section_num
            })
        
        # Если разделы не найдены, создаём один большой чанк
        if not chunks and text.strip():
            chunks.append({
                "id": f"{source_metadata['source_file']}_full",
                "text": text.strip(),
                "section_number": "Полный документ",
                "source_file": source_metadata["source_file"],
                "doc_type": "regulatory",
                "format": source_metadata.get("format", "unknown"),
                "page": 0
            })
        
        logger.info(f"   ✓ Создано {len(chunks)} чанков")
        return chunks
    
    @staticmethod
    def chunk_template(text: str, source_metadata: Dict) -> List[Dict]:
        """
        Разбиение шаблона документа на секции.
        """
        chunks = []
        
        # Разбиваем по заголовкам
        sections = re.split(r'\n(?=#{1,3}\s+|\n\d+[\.\)]\s+[А-Я])', text)
        
        section_num = 0
        for section in sections:
            section = section.strip()
            if not section or len(section) < 50:
                continue
            
            # Извлекаем название секции
            title_match = re.match(r'^(?:#{1,3}\s+)?(\d+[\.\)]\s+[А-Я][а-я]+|[А-Я]{3,})', section)
            section_name = title_match.group(1) if title_match else f"Раздел {section_num + 1}"
            
            section_num += 1
            chunks.append({
                "id": f"{source_metadata['source_file']}_section_{section_num}",
                "name": section_name.strip(),
                "content": section.strip(),
                "order": section_num,
                "source_file": source_metadata["source_file"],
                "doc_type": source_metadata.get("doc_type", "template"),
                "format": source_metadata.get("format", "unknown")
            })
        
        logger.info(f"   Создано {len(chunks)} секций шаблона")
        return chunks


class KnowledgeBaseIndexer:
    """Индексация в базы знаний"""
    
    def __init__(self):
        self.collection = get_chroma_collection()
        self.driver = get_neo4j_driver()
        self.embed_model = get_embedding_model()
    
    def index_regulatory_chunks(self, chunks: List[Dict]):
        """Индексация нормативных документов в ChromaDB"""
        if not chunks:
            logger.warning("Нет чанков для индексации")
            return
        
        logger.info(f"Индексация {len(chunks)} чанков в ChromaDB...")
        
        try:
            # Подготовка данных
            texts = [chunk["text"] for chunk in chunks]
            ids = [chunk["id"] for chunk in chunks]
            metadatas = [
                {
                    "section_number": chunk.get("section_number", ""),
                    "source_file": chunk["source_file"],
                    "doc_type": chunk["doc_type"],
                    "format": chunk.get("format", "unknown")
                }
                for chunk in chunks
            ]
            
            # Генерация эмбеддингов батчами
            batch_size = 16
            all_embeddings = []
            
            for i in range(0, len(texts), batch_size):
                batch = texts[i:i + batch_size]
                try:
                    batch_embeddings = self.embed_model.encode(batch).tolist()
                    all_embeddings.extend(batch_embeddings)
                    logger.info(f"   Обработано: {min(i + batch_size, len(texts))}/{len(texts)}")
                except Exception as e:
                    logger.error(f"   Ошибка генерации эмбеддингов: {e}")
                    continue
            
            # Добавление в ChromaDB
            if all_embeddings:
                self.collection.add(
                    documents=texts[:len(all_embeddings)],
                    embeddings=all_embeddings,
                    ids=ids[:len(all_embeddings)],
                    metadatas=metadatas[:len(all_embeddings)]
                )
                logger.info(f"   Добавлено {len(all_embeddings)} векторов")
            else:
                logger.warning("Не удалось сгенерировать эмбеддинги")
        
        except Exception as e:
            logger.error(f"Ошибка индексации в ChromaDB: {e}")
    
    def index_templates_to_neo4j(self, chunks: List[Dict]):
        """Индексация шаблонов в Neo4j"""
        if not chunks:
            logger.warning("Нет шаблонов для индексации")
            return
        
        logger.info(f"Индексация {len(chunks)} секций в Neo4j...")
        
        try:
            with self.driver.session() as session:
                for chunk in chunks:
                    # Создаём или находим узел шаблона
                    session.run("""
                        MERGE (t:Template {
                            doc_type: $doc_type,
                            source_file: $source
                        })
                        SET t.name = $name,
                            t.updated_at = datetime()
                    """,
                    doc_type=chunk["doc_type"],
                    source=chunk["source_file"],
                    name=chunk["source_file"].split('.')[0]
                    )
                    
                    # Создаём узел секции и связываем
                    session.run("""
                        MATCH (t:Template {
                            doc_type: $doc_type,
                            source_file: $source
                        })
                        CREATE (s:Section {
                            id: $id,
                            name: $name,
                            content: $content,
                            order: $order,
                            created_at: datetime()
                        })
                        MERGE (t)-[:CONTAINS {position: $order}]->(s)
                    """,
                    doc_type=chunk["doc_type"],
                    source=chunk["source_file"],
                    id=chunk["id"],
                    name=chunk["name"],
                    content=chunk["content"],
                    order=chunk["order"]
                    )
            
            logger.info(f"   Создано {len(chunks)} узлов в графе")
        
        except Exception as e:
            logger.error(f"Ошибка индексации в Neo4j: {e}")


def process_standards_directory(standards_dir: Path, indexer: KnowledgeBaseIndexer):
    """Обработка директории с нормативными документами"""
    if not standards_dir.exists():
        logger.warning(f"Директория не найдена: {standards_dir}")
        return
    
    logger.info("\n" + "=" * 60)
    logger.info("ИНДЕКСАЦИЯ НОРМАТИВНОЙ ДОКУМЕНТАЦИИ")
    logger.info("=" * 60)
    
    parser = DocumentParser()
    chunker = TextChunker()
    all_chunks = []
    
    # Обработка файлов различных форматов
    file_patterns = {
        "*.pdf": parser.parse_pdf,
        "*.docx": parser.parse_docx,
        "*.odt": parser.parse_odt,
        "*.txt": parser.parse_text,
        "*.md": parser.parse_text
    }
    
    for pattern, parse_func in file_patterns.items():
        for file_path in standards_dir.glob(pattern):
            logger.info(f"\nОбработка: {file_path.name}")
            
            # Извлечение текста
            text, metadata = parse_func(file_path)
            
            if not text:
                logger.warning(f"Пустой или не распаршен файл: {file_path.name}")
                continue
            
            # Разбиение на чанки
            chunks = chunker.chunk_regulatory_document(text, metadata)
            all_chunks.extend(chunks)
    
    # Индексация
    if all_chunks:
        indexer.index_regulatory_chunks(all_chunks)
        logger.info(f"\nИндексация нормативки завершена: {len(all_chunks)} разделов")
    else:
        logger.warning("\nНормативные документы не найдены или не распаршены")


def process_templates_directory(templates_dir: Path, indexer: KnowledgeBaseIndexer):
    """Обработка директории с шаблонами"""
    if not templates_dir.exists():
        logger.warning(f"Директория не найдена: {templates_dir}")
        return
    
    logger.info("\n" + "=" * 60)
    logger.info("ИНДЕКСАЦИЯ ШАБЛОНОВ ДОКУМЕНТОВ")
    logger.info("=" * 60)
    
    parser = DocumentParser()
    chunker = TextChunker()
    all_templates = []
    
    # Обработка DOCX шаблонов
    for file_path in templates_dir.glob("*.docx"):
        logger.info(f"\nОбработка шаблона: {file_path.name}")
        
        text, metadata = parser.parse_docx(file_path)
        
        if not text:
            logger.warning(f"Пустой или не распаршен шаблон: {file_path.name}")
            continue
        
        # Определяем тип документа из имени файла
        doc_type = file_path.stem.replace('_', ' ').replace('-', ' ')
        metadata["doc_type"] = doc_type
        
        # Разбиение на секции
        chunks = chunker.chunk_template(text, metadata)
        all_templates.extend(chunks)
    
    # Индексация
    if all_templates:
        indexer.index_templates_to_neo4j(all_templates)
        logger.info(f"\nИндексация шаблонов завершена: {len(all_templates)} секций")
    else:
        logger.warning("\nШаблоны не найдены или не распаршены")


def main():
    """Основная функция"""
    logger.info("\n" + "=" * 60)
    logger.info("ЗАПУСК ИНДЕКСАЦИИ БАЗЫ ЗНАНИЙ")
    logger.info("=" * 60)
    
    # Создание директорий для логов
    Path("logs").mkdir(exist_ok=True)
    
    # Инициализация индексера
    try:
        indexer = KnowledgeBaseIndexer()
        logger.info("Подключение к базам данных успешно")
    except Exception as e:
        logger.error(f"Ошибка подключения к БД: {e}")
        sys.exit(1)
    
    # Обработка нормативных документов
    standards_dir = Path("./data/documents/standards")
    process_standards_directory(standards_dir, indexer)
    
    # Обработка шаблонов
    templates_dir = Path("./data/documents/templates")
    process_templates_directory(templates_dir, indexer)
    
    # Итоговая статистика
    logger.info("\n" + "=" * 60)
    logger.info("ИНДЕКСАЦИЯ ЗАВЕРШЕНА")
    logger.info("=" * 60)
    
    # Проверка количества документов
    try:
        count = indexer.collection.count()
        logger.info(f"Всего документов в ChromaDB: {count}")
    except Exception as e:
        logger.warning(f"Не удалось получить количество документов: {e}")


if __name__ == "__main__":
    main()